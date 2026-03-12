#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
面试评估与规划 Word 文档生成器

输入 JSON 数据结构：
{
  "candidate_name": "候选人姓名",
  "position": "岗位名称",
  "basic_info": [
    {"label": "姓名", "value": "XXX"},
    {"label": "年龄", "value": "29岁"},
    ...
  ],
  "advantages": [
    {"title": "优势标题", "detail": "具体描述"},
    ...
  ],
  "disadvantages": [
    {"title": "劣势标题", "detail": "具体描述"},
    ...
  ],
  "match_note": "匹配度说明文字（可选，如侧重方向说明）",
  "match_table": [
    {"dimension": "评估维度", "requirement": "JD要求", "match": "候选人匹配", "score": "★★★★☆"},
    ...
  ],
  "match_summary": "综合评价文字",
  "match_detail": "匹配度详细说明",
  "interview_plan": {
    "total_time": "约90-120分钟，建议分两轮",
    "round1": {
      "title": "第一轮：技术面试（60-70分钟）",
      "parts": [
        {
          "title": "Part 1：项目经验深挖（15分钟）",
          "intro": "可选的 part 说明文字",
          "questions": [
            {"q": "问题内容", "a": "参考答案/考察点"}
          ]
        }
      ]
    },
    "round2": {
      "title": "第二轮：综合面试（30-40分钟）",
      "questions": [
        {"q": "问题内容", "a": "参考答案/考察点"}
      ]
    }
  },
  "score_table": {
    "dimensions": [
      {"dimension": "评估维度", "weight": "25%", "criteria": "Q1-Q3表现"}
    ],
    "thresholds": [
      "总分 ≥ 70分：建议录用",
      "总分 60-69分：条件录用（明确培养计划）",
      "总分 < 60分：不建议录用"
    ]
  },
  "summary": {
    "portrait": "候选人整体画像描述",
    "focus_points": [
      "面试重点关注事项1",
      "面试重点关注事项2"
    ]
  }
}
"""

import argparse
import json
import sys

def generate_docx(data, output_path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title = doc.add_heading('', level=0)
    run = title.add_run(f'{data["candidate_name"]}_{data["position"]}_面试评估与规划')
    run.font.size = Pt(22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === 基本信息 ===
    doc.add_heading('一、候选人基本信息', level=1)
    info = data['basic_info']
    table = doc.add_table(rows=len(info), cols=2, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, item in enumerate(info):
        table.cell(i, 0).text = item['label']
        table.cell(i, 1).text = item['value']
        for cell in [table.cell(i, 0), table.cell(i, 1)]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10.5)
    doc.add_paragraph('')

    # === 优势 ===
    doc.add_heading('二、候选人优势', level=1)
    for adv in data['advantages']:
        p = doc.add_paragraph()
        run = p.add_run(adv['title'])
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(adv['detail'])

    # === 劣势 ===
    doc.add_heading('三、候选人劣势', level=1)
    for dis in data['disadvantages']:
        p = doc.add_paragraph()
        run = p.add_run(dis['title'])
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(dis['detail'])

    # === 匹配度 ===
    doc.add_heading('四、岗位匹配度评估', level=1)

    if data.get('match_note'):
        doc.add_paragraph(data['match_note'])

    mt = data['match_table']
    table2 = doc.add_table(rows=len(mt) + 1, cols=4, style='Light Grid Accent 1')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['评估维度', 'JD要求', '候选人匹配', '评分']
    for i, h in enumerate(headers):
        table2.cell(0, i).text = h
        for p in table2.cell(0, i).paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(mt):
        table2.cell(i + 1, 0).text = row['dimension']
        table2.cell(i + 1, 1).text = row['requirement']
        table2.cell(i + 1, 2).text = row['match']
        table2.cell(i + 1, 3).text = row['score']

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run(data['match_summary'])
    run.bold = True
    run.font.size = Pt(11)
    if data.get('match_detail'):
        doc.add_paragraph(data['match_detail'])

    # === 面试流程 ===
    doc.add_heading('五、面试流程规划', level=1)
    plan = data['interview_plan']
    doc.add_paragraph(plan['total_time'])

    # 第一轮
    doc.add_heading(plan['round1']['title'], level=2)
    for part in plan['round1']['parts']:
        doc.add_heading(part['title'], level=3)
        if part.get('intro'):
            doc.add_paragraph(part['intro'])
        for item in part['questions']:
            p = doc.add_paragraph()
            run = p.add_run(item['q'])
            run.bold = True
            run.font.size = Pt(10.5)
            p2 = doc.add_paragraph()
            run2 = p2.add_run('参考答案：')
            run2.bold = True
            run2.font.color.rgb = RGBColor(0, 102, 153)
            p2.add_run(item['a'])
            doc.add_paragraph('')

    # 第二轮
    doc.add_heading(plan['round2']['title'], level=2)
    for item in plan['round2']['questions']:
        p = doc.add_paragraph()
        run = p.add_run(item['q'])
        run.bold = True
        run.font.size = Pt(10.5)
        p2 = doc.add_paragraph()
        run2 = p2.add_run('参考答案/考察点：')
        run2.bold = True
        run2.font.color.rgb = RGBColor(0, 102, 153)
        p2.add_run(item['a'])
        doc.add_paragraph('')

    # === 评分建议 ===
    doc.add_heading('六、面试评分建议表', level=1)
    dims = data['score_table']['dimensions']
    table3 = doc.add_table(rows=len(dims) + 1, cols=3, style='Light Grid Accent 1')
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['评估维度', '权重', '评分标准']):
        table3.cell(0, i).text = h
        for p in table3.cell(0, i).paragraphs:
            for r in p.runs:
                r.bold = True
    for i, d in enumerate(dims):
        table3.cell(i + 1, 0).text = d['dimension']
        table3.cell(i + 1, 1).text = d['weight']
        table3.cell(i + 1, 2).text = d['criteria']

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('录用建议阈值：')
    run.bold = True
    for t in data['score_table']['thresholds']:
        doc.add_paragraph(t, style='List Bullet')

    # === 总结 ===
    doc.add_heading('七、总结与建议', level=1)
    p = doc.add_paragraph()
    run = p.add_run('候选人整体画像：')
    run.bold = True
    p.add_run(data['summary']['portrait'])
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('面试重点关注：')
    run.bold = True
    for fp in data['summary']['focus_points']:
        doc.add_paragraph(fp, style='List Bullet')

    doc.save(output_path)
    print(f'文档已保存至: {output_path}')


def main():
    parser = argparse.ArgumentParser(description='生成面试评估与规划 Word 文档')
    parser.add_argument('--data', required=True, help='JSON 数据文件路径')
    parser.add_argument('--output', required=True, help='输出 docx 文件路径')
    args = parser.parse_args()

    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)

    generate_docx(data, args.output)


if __name__ == '__main__':
    main()
